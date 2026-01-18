from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import pytz
from dotenv import dotenv_values

# Load .env
env_path = ".env"
config = dotenv_values(env_path)
area_path = config["AREA_PATH"]
iteration_path = config["ITERATION_PATH"]
project = config["ADO_PROJECT"]
end_date = config["END_DATE"]

doc = Document()
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from zoneinfo import ZoneInfo
from datetime import datetime

# Convert END_DATE string to datetime and format for EST
end_date_dt = datetime.strptime(end_date, "%Y-%m-%d")
end_date_est = end_date_dt.astimezone(ZoneInfo("America/New_York")).strftime("%B %d, %Y")

# Access header
section = doc.sections[0]
header = section.header

# Use the first paragraph (or create one)
hdr_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
hdr_para.clear() if hasattr(hdr_para, 'clear') else None

# Compose inline content: logo left, then tab, then right-aligned text
logo_run = hdr_para.add_run()
logo_run.add_picture("flwins_logo.png", width=Inches(1.0))

# Tabs: one moves to center, two to right
text = "\t\tFL WINS UAT Test Report"
text_run = hdr_para.add_run(text)
text_run.font.bold = True

# Add page number on the bottom right of the footer
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Access footer
footer = section.footer
footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

# Insert PAGE field
run = footer_para.add_run()
fldChar1 = OxmlElement('w:fldChar')       # create field character element
fldChar1.set(qn('w:fldCharType'), 'begin')

instrText = OxmlElement('w:instrText')    # instruction
instrText.set(qn('xml:space'), 'preserve')
instrText.text = "PAGE"

fldChar2 = OxmlElement('w:fldChar')       # create end field character
fldChar2.set(qn('w:fldCharType'), 'end')

# Append all the XML elements to the run
run._r.append(fldChar1)
run._r.append(instrText)
run._r.append(fldChar2)

# Add and style the title
para = doc.add_paragraph()
run = para.add_run(project+" UAT Test Report")
run.bold = True
run.font.size = Pt(24)
para.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add EST date-time stamp
est = pytz.timezone("US/Eastern")
timestamp = datetime.now(est).strftime("%B %d, %Y %I:%M %p EST")

ts_para = doc.add_paragraph()
ts_run = ts_para.add_run(timestamp)
ts_run.font.size = Pt(12)
ts_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Area & Iteration (side by side, in one line)
info_text = f"Area Path: {area_path}    |    Iteration Path: {iteration_path}"
info_para = doc.add_paragraph()
info_run = info_para.add_run(info_text)
info_run.font.size = Pt(12)
info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

##### TOTAL BUG BURN DOWN CHART #####

import subprocess
from docx.shared import Inches
import time
import os

# Step 1: Run the script to generate the chart
subprocess.run(["python", "total_bug_burndown.py"], check=True)

# Step 2: Wait for file to exist (or retry briefly)
chart_path = "charts/burndown_chart.jpg"
for _ in range(10):  # wait up to ~5 seconds
    if os.path.exists(chart_path):
        break
    time.sleep(0.5)
else:
    raise FileNotFoundError(f"Expected chart not found at {chart_path}")

# Step 3: Insert image into the doc
chart_para = doc.add_paragraph()
run = chart_para.add_run()
run.add_picture(chart_path, width=Inches(6.0))  # adjust width if needed
chart_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

#### Bug report chart #####

# Step: Run daily_bug_report.py to generate latest charts
import subprocess, os, time
subprocess.run(['python', 'daily_bug_report.py'], check=True)

# Step: Insert status_summary.jpg into the document when ready
status_chart = 'charts/status_summary.jpg'
for _ in range(20):  # wait up to 10 seconds
    if os.path.exists(status_chart):
        doc.add_picture(status_chart, width=Inches(6))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        break
    time.sleep(0.5)

# Step: Insert severity_summary.jpg into the document when ready
severity_chart = 'charts/severity_summary.jpg'
for _ in range(20):  # wait up to 10 seconds
    if os.path.exists(severity_chart):
        doc.add_picture(severity_chart, width=Inches(6))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        break
    time.sleep(0.5)

# Step: Optionally insert severity_4_5_bug_table_final.jpg
bug_table = 'charts/severity_4_5_bug_table.jpg'
time.sleep(2)  # allow time to generate
if os.path.exists(bug_table):
    doc.add_picture(bug_table, width=Inches(6))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

#### Projection chart #####

# # Step: Run test_case_burndown copy.py to generate latest charts
# import subprocess, os, time
# subprocess.run(['python', 'test_case_burndown copy.py'], check=True)

# # Step: Optionally insert severity_4_5_bug_table_final.jpg
# test_chart = 'charts/test_case_burndown_with_plan.jpg'
# time.sleep(240)  # allow time to generate
# if os.path.exists(bug_table):
#     doc.add_picture(test_chart, width=Inches(6))
#     last_paragraph = doc.paragraphs[-1]
#     last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


# Generate file timestamp
file_timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")

# Save the final report
doc.save(f"reports/uat_test_report_{file_timestamp}.docx")
